from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from build_common import *

TITLE = ("Lineage-Aware Hierarchical Deep Learning for Imbalanced Multi-Class "
     "Classification of Peripheral Blood Cells")


def front_matter(doc):
    C = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3):
        para(doc)
    para(doc, TITLE, align=C, bold=True, size=18)
    para(doc)
    para(doc, "Final Thesis", align=C, size=14)
    para(doc)
    para(doc, "In Partial Fulfillment", align=C, size=12)
    para(doc, "of the Requirements for the Degree of", align=C, size=12)
    para(doc, "Master in Computer Science", align=C, bold=True, size=14)
    for _ in range(4):
        para(doc)
    for label, value in (("Student Name", "Arthiga Karthigesu"),
                         ("Student ID", "[insert student ID]"),
                         ("Supervisor", "[insert supervisor name]")):
        p = doc.add_paragraph()
        p.alignment = C
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}  :  {value}")
        r.font.name = FONT
        r.font.size = Pt(12)
    pagebreak(doc)

    # ---------------- Abstract ----------------
    h1(doc, "Abstract")
    para(doc,
         "Automated classification of peripheral blood cells is a well-studied "
              "application of deep learning, yet two properties of the clinical problem "
              "are routinely discarded. The eighteen cell types encountered in a "
              "differential count are not unrelated labels: they belong to three "
              "haematopoietic lineages, and within the myeloid branch they lie on a "
              "maturation continuum. The class distribution is also severely long-tailed, "
              "so the rare types carrying the greatest diagnostic weight are learned "
              "least well. This dissertation investigates whether making a classifier "
              "lineage-aware, and training it under an imbalance-robust objective, "
              "improves fine-grained classification of rare blood cell types relative to "
              "a flat transfer-learning baseline.")
    para(doc,
         "Using the MLL23 dataset of 41,621 expert-annotated images across 18 "
              "classes at an imbalance ratio of 261:1, a shared pretrained backbone feeds "
              "lineage and fine cell-type heads coupled by a class-balanced hierarchical "
              "loss combining effective-number reweighting, focal modulation and a "
              "consistency term. Four configurations were trained under identical "
              "partitions and backbones at three seeds each, and compared by paired "
              "testing with reported effect sizes.")
    para(doc,
         "The hierarchical model reached a test macro F1 of 0.879 against 0.874 for "
              "the flat baseline, and a minority-class macro F1 of 0.840 against 0.823, "
              "improving rare classes by more than the overall average. Aggregate "
              "differences were not significant at three seeds, but the model "
              "significantly reduced cross-lineage error, the clinically severe category "
              "(p = 0.020), confirming the predicted mechanism; per-class analysis "
              "located the gain almost entirely in the two rarest lymphoid types. "
              "Removing the imbalance term cost 0.060 minority macro F1, an order of "
              "magnitude more than the hierarchy contributed, identifying the objective "
              "as the load-bearing component. Lineage accuracy saturated at 98.9%, "
              "explaining the modest architectural effect and indicating that the "
              "auxiliary signal should instead target maturation stage.")
    para(doc)
    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r2 = p.add_run("hierarchical classification; class imbalance; peripheral blood cells; "
         "class-balanced loss; explainable AI; deep learning")
    r2.font.name = "Arial"
    r2.font.size = Pt(11)
    pagebreak(doc)

    # ---------------- Acknowledgements ----------------
    h1(doc, "Acknowledgements")
    para(doc,
         "I would like to thank my supervisor for their guidance throughout this "
              "project, and in particular for the encouragement to treat negative and "
              "inconclusive findings as results in their own right, not failures to be "
              "tidied away.")
    para(doc,
         "I am grateful to the Munich Leukemia Laboratory and the authors of the "
              "MLL23 data descriptor for releasing a large, expert-annotated corpus under "
              "an open licence. Work of this kind is only possible because practitioners "
              "choose to make carefully curated clinical data publicly available.")
    para(doc, "Finally, I thank my family for their patience and support.")
    pagebreak(doc)

    # ---------------- Contents ----------------
    h1(doc, "Contents")
    toc = [
        ("Abstract", "ii"), ("Acknowledgements", "iii"), ("Contents", "iv"),
        ("List of Tables", "v"), ("List of Figures", "vi"),
        ("List of Acronyms", "vii"),
        ("Chapter 1  Introduction", "1"),
        ("     1.1  Background", "1"),
        ("     1.2  Problem Statement", "2"),
        ("     1.3  Research Question and Objectives", "3"),
        ("     1.4  Expected Outcomes", "4"),
        ("Chapter 2  Literature Review", "5"),
        ("     2.1  Comprehensive Overview of the Existing Literature", "5"),
        ("     2.2  Critical Analysis of Existing Studies", "13"),
        ("Chapter 3  Methodology", "17"),
        ("     3.1  Data Collection and Preprocessing", "18"),
        ("     3.2  ML/AI Model Development", "21"),
        ("     3.3  Evaluation of the Proposed System", "24"),
        ("     3.4  Experimental Design and Statistical Protocol", "26"),
        ("Chapter 4  Experimental Results", "28"),
        ("     4.1  Experimental Setup", "28"),
        ("     4.2  Dataset Description", "29"),
        ("     4.3  Results", "30"),
        ("     4.4  Comparison with Baseline Methods", "36"),
        ("Chapter 5  Conclusion and Future Work", "40"),
        ("     5.1  Conclusion", "40"),
        ("     5.2  Future Work", "42"),
        ("References", "44"),
        ("Appendix A. Experimental Configuration Reference", "48"),
        ("Appendix B. Per-Class Results", "49"),
    ]
    for label, pg in toc:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        r.font.name = FONT
        r.font.size = Pt(12)
        if not label.startswith(" "):
            r.bold = True
        r2 = p.add_run("\t" + pg)
        r2.font.name = FONT
        r2.font.size = Pt(12)
    pagebreak(doc)

    # ---------------- List of Tables ----------------
    h1(doc, "List of Tables")
    for lbl in [
        "Table 2.1  Critical analysis and summary of the existing studies",
        "Table 3.1  The 18 MLL23 classes, lineage membership and partition sizes",
        "Table 3.2  Class-conditional augmentation parameters",
        "Table 3.3  Candidate backbone architectures",
        "Table 3.4  Ablation switches implemented on the loss object",
        "Table 3.5  The four distinct configurations and five reported experiments",
        "Table 4.1  Training hyperparameters, held identical across all arms",
        "Table 4.2  Test-set performance by experimental arm",
        "Table 4.3  Paired comparisons across three seeds",
        "Table 4.4  Hierarchical decomposition of prediction outcomes",
        "Table 4.5  Per-class test F1 by arm",
        "Table A1  Configuration fields and Phase 2 values",
        "Table A2  Computational cost of the Phase 2 matrix",
        "Table B1  Per-class precision and recall on the test partition",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(lbl)
        r.font.name = FONT
        r.font.size = Pt(12)
    pagebreak(doc)

    # ---------------- List of Figures ----------------
    h1(doc, "List of Figures")
    for lbl in [
        "Figure 3.1  Architecture of the proposed system",
        "Figure 3.2  The two-level label hierarchy",
        "Figure 3.3  One representative image per class after preprocessing",
        "Figure 3.4  Class distribution on linear and logarithmic axes",
        "Figure 3.5  The three augmentation policies applied to one image",
        "Figure 3.6  Stain normalisation, one image per lineage",
        "Figure 3.7  UMAP projection of frozen backbone features by lineage",
        "Figure 4.1  Overall and minority-class F1 by experimental arm",
        "Figure 4.2  Cross-lineage error for each seed, baseline against proposed",
        "Figure 4.3  Per-class F1, proposed model against the flat baseline",
        "Figure 4.4  How far along the maturation chain each error lands",
        "Figure 4.5  Grad-CAM saliency for the fine classification head",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(lbl)
        r.font.name = FONT
        r.font.size = Pt(12)
    pagebreak(doc)

    # ---------------- List of Acronyms ----------------
    h1(doc, "List of Acronyms")
    table(doc, ["Term", "Definition"], [
        ("AML", "Acute Myeloid Leukaemia"),
        ("AMP", "Automatic Mixed Precision"),
        ("bf16", "Brain Floating Point, 16-bit"),
        ("CB", "Class-Balanced loss"),
        ("CNN", "Convolutional Neural Network"),
        ("cRT", "Classifier Re-Training, decoupled training"),
        ("EMA", "Exponential Moving Average"),
        ("Grad-CAM", "Gradient-weighted Class Activation Mapping"),
        ("KL", "Kullback-Leibler divergence"),
        ("LDAM", "Label-Distribution-Aware Margin loss"),
        ("MLL23", "Munich Leukemia Laboratory 2023 dataset"),
        ("TTA", "Test-Time Augmentation"),
        ("UMAP", "Uniform Manifold Approximation and Projection"),
        ("ViT", "Vision Transformer"),
        ("XAI", "Explainable Artificial Intelligence"),
    ], widths=[1.5, 4.2], font_size=11)
    pagebreak(doc)


def chapter1(doc):
    h1(doc, "Chapter 1  Introduction")
    para(doc,
         "This chapter establishes the clinical and technical context of the study, "
              "states the problem it addresses, and sets out the research question, "
              "objectives and expected outcomes that structure the remainder of the "
              "dissertation.")

    h2(doc, "1.1  Background")
    para(doc,
         "The microscopic examination of a peripheral blood smear remains a "
              "foundational procedure in haematological diagnosis. In a differential "
              "count, a trained cytomorphologist inspects individual nucleated cells and "
              "assigns each to a cell type on the basis of nuclear shape, chromatin "
              "texture, cytoplasmic volume and colour, and the presence and character of "
              "granulation. The resulting distribution of cell types informs the "
              "diagnosis and monitoring of conditions ranging from reactive infection to "
              "acute leukaemia [1], [2].")
    para(doc,
         "The procedure has two well-documented weaknesses. It is slow and "
              "repetitive, requiring sustained attention across hundreds of cells per "
              "slide, and it is subject to appreciable inter-observer variability, since "
              "the morphological boundaries between adjacent cell types shade into one "
              "another instead of falling into clean categories [3]. Both weaknesses are "
              "most pronounced for the rare and immature cell types, which a given "
              "observer encounters infrequently and whose appearance overlaps "
              "substantially with their neighbours on the maturation continuum.")
    para(doc,
         "Deep learning has been applied to this task with considerable success. "
              "Convolutional neural networks trained on large annotated corpora have "
              "reached expert-level agreement on blast cell recognition [4] and on "
              "bone-marrow cell morphology across a wide class vocabulary [2]. The "
              "publication of several large, expert-annotated single-cell datasets has "
              "accelerated this work by providing common benchmarks [1], [5], [6]. "
              "Transfer learning from ImageNet-pretrained backbones is now the standard "
              "approach, and reported accuracies on the more balanced benchmarks "
              "routinely exceed 95%.")
    para(doc,
         "Two properties of the clinical problem are nonetheless discarded by most "
              "of these systems. The first is the structure of the label space. The cell "
              "types encountered in a differential count are not unrelated categories: "
              "they belong to three haematopoietic lineages, and within the myeloid "
              "branch they occupy successive stages of a single developmental process. A "
              "classifier that treats the eighteen types as a flat vocabulary discards "
              "this biology entirely. The second is the shape of the class distribution. "
              "Blood cell datasets that reflect genuine clinical prevalence are severely "
              "long-tailed, and the rare types are frequently the diagnostically decisive "
              "ones.")

    h2(doc, "1.2  Problem Statement")
    para(doc,
         "A model trained conventionally on a long-tailed corpus optimises an "
              "objective dominated by the majority classes. Where one class contributes "
              "8,606 images and another 33, the gradient contributed by the rare class is "
              "negligible, and a classifier that abandons it entirely forfeits less than "
              "0.1% of overall accuracy. The consequence is a system that performs "
              "impressively on the headline metric while failing on exactly the cells "
              "that motivated its construction. Reporting plain accuracy on such data is "
              "worse than uninformative; it actively misleads [7], [8].")
    para(doc,
         "The conventional remedy is to resample the training distribution until it "
              "is balanced. This is unsatisfactory in the present context for two "
              "reasons. Oversampling a class of 33 images cannot manufacture "
              "morphological variety that the data does not contain, and risks "
              "memorisation of the few available examples; undersampling the majority "
              "discards the large volume of data that makes transfer learning effective "
              "in the first place. More fundamentally, the imbalance reflects genuine "
              "physiological prevalence, and is the phenomenon under study, not a defect "
              "to be cleaned up before study begins.")
    para(doc,
         "A second limitation concerns the treatment of errors. A flat classifier "
              "assigns equal cost to every misclassification, yet the clinical "
              "consequences differ sharply. Confusing a myelocyte with a metamyelocyte, "
              "two adjacent stages of neutrophil maturation, is a mild error that "
              "subsequent human review would likely tolerate. Assigning a lymphocyte to "
              "the myeloid lineage is a categorically different failure with direct "
              "diagnostic consequences. Standard metrics cannot express this distinction, "
              "and standard objectives provide no mechanism to prefer the milder error.")
    para(doc,
         "Finally, systems intended for clinical deployment must be interpretable. A "
              "model that achieves high accuracy by attending to staining artefacts or to "
              "neighbouring red cells instead of the target cell will not generalise to a "
              "new laboratory, and cannot be trusted by the practitioner expected to act "
              "on its output [9], [10]. Saliency analysis is built into the design, not "
              "added at the end as illustration.")
    para(doc,
         "These three concerns, the discarded label hierarchy, the long-tailed "
              "distribution, and the need for interpretable evidence, have each been "
              "addressed in isolation in the literature. They have not been addressed "
              "together on a large, genuinely imbalanced peripheral blood corpus, and it "
              "is that combination this dissertation examines.")

    h2(doc, "1.3  Research Question and Objectives")
    para(doc, "This study is guided by the following research question.")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        "Does a lineage-aware hierarchical classifier, combined with an "
             "imbalance-robust training objective, improve fine-grained classification "
             "of single peripheral blood cell images, especially for rare cell types, "
             "compared with a flat transfer-learning classifier on the same data?")
    r.italic = True
    r.font.name = FONT
    r.font.size = Pt(12)
    para(doc, "The question is addressed through five objectives.")
    bullets(doc, [
        "Acquire, verify and characterise a large, expert-annotated peripheral "
             "blood single-cell dataset exhibiting genuine clinical class imbalance, and "
             "establish a deterministic partition that preserves rare classes in every "
             "split.",
        "Design and implement a hierarchical classification architecture in which a "
             "shared pretrained backbone supports both a lineage head and a fine "
             "cell-type head, trained jointly under a single composite objective.",
        "Formulate an imbalance-robust loss combining effective-number class "
             "reweighting, focal modulation and a hierarchical consistency term, with "
             "each component independently removable for ablation.",
        "Evaluate the proposed model against a flat baseline and against ablations "
             "removing the hierarchy and the imbalance treatment, under identical "
             "partitions and backbones, using macro-averaged and minority-class metrics "
             "together with a hierarchical decomposition of errors.",
        "Assess whether the decisions of the model are grounded in morphologically "
             "meaningful image regions using gradient-based saliency analysis.",
    ], numbered=True)

    h2(doc, "1.4  Expected Outcomes")
    para(doc,
         "The primary expected outcome is quantitative evidence on whether lineage "
              "awareness improves rare-class performance. The design anticipates one "
              "specific pattern, not a general lift: if the hierarchy contributes as "
              "theorised, the gain should be larger on the minority classes than on the "
              "corpus as a whole, and errors should shift from cross-lineage to "
              "within-lineage even where total error moves little.")
    para(doc,
         "The ablation design is expected to apportion credit between the two "
              "mechanisms under test. Because the hierarchy and the imbalance objective "
              "are removable independently, the study can report not only whether the "
              "combined system outperforms the baseline but which component is "
              "responsible, which is the more useful finding for subsequent work.")
    para(doc,
         "The saliency analysis is expected to establish whether the network attends "
              "to nucleus, chromatin and cytoplasm instead of background or neighbouring "
              "red cells, providing qualitative support for or against the quantitative "
              "results.")
    para(doc,
         "Finally, the study is expected to produce a reusable, configuration-driven "
              "experimental framework in which every arm is a parameter setting on one "
              "implementation instead of a script of its own, so that the comparisons "
              "remain valid as further variants are added.")
    para(doc,
         "It is worth stating what the study does not expect to deliver. It does not "
              "aim to produce a deployable diagnostic system, since a single-centre "
              "corpus cannot establish the cross-institution robustness that deployment "
              "would require. Nor does it aim to establish which backbone architecture is "
              "best suited to the task, which would demand a full-length comparison "
              "beyond the available compute budget. The contribution is a controlled test "
              "of two specific mechanisms, evaluated under conditions that make the "
              "comparison between them interpretable.")

    h2(doc, "1.5  Structure of the Dissertation")
    para(doc,
         "Chapter 2 reviews the literature across the four domains the problem draws "
              "on, analyses the existing studies against a common set of parameters, and "
              "identifies the gap the study addresses. Chapter 3 sets out the "
              "methodology: the dataset and its label hierarchy, the partitioning and "
              "preprocessing procedure, the two-head architecture, the composite training "
              "objective and its ablation switches, the evaluation metrics including the "
              "hierarchical decomposition of errors, and the experimental design under "
              "which the arms are compared.")
    para(doc,
         "Chapter 4 reports the experimental setup and results. It presents the four "
              "arms at aggregate level, analyses them per class, examines the saliency "
              "evidence, compares the outcome against the flat baseline and against the "
              "literature, and records a methodological correction that invalidated an "
              "earlier experimental matrix. Chapter 5 draws conclusions against the "
              "research question and objectives, and sets out the further work the "
              "findings indicate. Appendix A documents the experimental configuration in "
              "full, and Appendix B reports per-class precision and recall for every arm.")
    pagebreak(doc)
