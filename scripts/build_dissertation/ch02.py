from docx.shared import Pt, Inches
from build_common import *


def chapter2(doc):
    h1(doc, "Chapter 2  Literature Review")
    para(doc,
         "The automated classification of single peripheral blood cells draws on "
              "four distinct but intersecting research domains: deep learning "
              "architectures applied to cytological imagery, hierarchical classification "
              "paradigms that exploit known label structure, algorithm-level strategies "
              "for learning from severely imbalanced data, and explainable artificial "
              "intelligence for clinical viability. Each has developed substantially over "
              "the past decade. Their synthesis, particularly within peripheral blood "
              "diagnostics, remains underdeveloped, and it is that absence which defines "
              "the contribution of the present study.")
    para(doc,
         "This review is grounded in primary research published between 2021 and "
              "2026, with preference for Q1-ranked journals in computational medicine, "
              "biomedical engineering and computer vision, supplemented by a small number "
              "of foundational earlier works whose methods remain in active use. The "
              "search strategy used the keywords peripheral blood cell classification, "
              "leukaemia deep learning, hierarchical classification medical imaging, "
              "class imbalance loss function, explainable AI haematology, and blood smear "
              "convolutional neural network, cross-referenced with haematopoiesis "
              "taxonomy terms. Candidate papers were retained where they reported "
              "quantitative evaluation on cytological image data, or where they "
              "introduced a method subsequently adopted in that setting. For each "
              "retained study the problem addressed, the methodology, the dataset, the "
              "reported results and the stated limitations were recorded, and it is that "
              "record which structures the critical analysis in Section 2.2.")

    h2(doc, "2.1  Comprehensive Overview of the Existing Literature")

    h3(doc, "2.1.1  From Handcrafted Features to End-to-End Learning")
    para(doc,
         "Early automated blood cell analysis relied on explicit segmentation "
              "followed by handcrafted feature extraction. A typical pipeline isolated "
              "the nucleus and cytoplasm by colour-space thresholding, computed "
              "morphometric descriptors such as area, circularity, nucleus-to-cytoplasm "
              "ratio and texture statistics, and passed these to a conventional "
              "classifier such as a support vector machine or random forest. The approach "
              "is interpretable by construction, since every feature has a stated meaning "
              "that a haematologist can evaluate, and it is economical with training data "
              "because the feature space is small. Its ceiling, however, is set by the "
              "expressiveness of the chosen descriptors, and morphological distinctions "
              "that no one thought to encode are simply invisible to the classifier.")
    para(doc,
         "Rahimzadeh and colleagues [5] published the Raabin-WBC dataset of "
              "approximately 40,000 peripheral white blood cell images acquired with two "
              "camera systems and two microscope platforms. Their experiments "
              "demonstrated that the generalisation of machine learning models is "
              "profoundly affected by acquisition diversity, a challenge handcrafted "
              "pipelines cannot address because their descriptors are camera-agnostic by "
              "design and so blind to the systematic colour and focus shifts that "
              "different hardware introduces. This finding has direct implications for "
              "deployment, where imaging hardware is rarely standardised across "
              "institutions and where a model validated on one microscope may be applied "
              "on another.")
    para(doc,
         "Kouzehkanan and colleagues [6] proposed a segmentation-first, "
              "feature-extraction-second approach that outperformed several pretrained "
              "convolutional networks in cross-dataset generalisation, because it "
              "combined principled colour-space segmentation with robust shape features "
              "that transfer across acquisition conditions. Their result is a useful "
              "corrective to the assumption that end-to-end learning dominates in every "
              "respect: where the training corpus is narrow, explicit priors can "
              "generalise better than learned ones. It simultaneously illustrated the "
              "performance ceiling such pipelines impose, since the same rigidity that "
              "confers robustness prevents the model from exploiting subtle texture cues "
              "that distinguish adjacent maturation stages. Acevedo and colleagues [26] "
              "provided an early demonstration that convolutional networks could exceed "
              "this ceiling on peripheral blood images, establishing the transition that "
              "the subsequent literature consolidated.")

    h3(doc, "2.1.2  Convolutional Architectures for Blood Cell Classification")
    para(doc,
         "Convolutional neural networks changed the pipeline by allowing the model "
              "to discover hierarchical feature representations directly from pixels, "
              "removing the segmentation step and the descriptor design step "
              "simultaneously. Matek and colleagues [2] trained convolutional networks on "
              "171,374 annotated single-cell images from 945 haematologically diverse "
              "patients, demonstrating that such networks identify a wide range of "
              "diagnostically relevant morphologies, including subtle blast appearances, "
              "with precision and recall approaching expert haematologists. The dataset "
              "was the largest expert-annotated pool of bone marrow cytology published at "
              "the time and remains a reference benchmark. An earlier study by the same "
              "group established human-level agreement on blast recognition in acute "
              "myeloid leukaemia [4], which is the result most often cited as evidence "
              "that the task is tractable at all.")
    para(doc,
         "Transfer learning from ImageNet-pretrained backbones subsequently became "
              "the dominant paradigm, and for good reason: medical image corpora are "
              "typically one to two orders of magnitude smaller than general-purpose "
              "vision benchmarks, and pretrained low-level features transfer well because "
              "edges, textures and blobs are not domain-specific. Alam and Islam [11] "
              "fine-tuned DenseNet-121 on a four-class dataset of 12,444 images, "
              "reporting 98.84% accuracy and 99.61% specificity, and showed through "
              "systematic ablation that batch size materially affects performance in this "
              "regime, with smaller batches acting as an implicit regulariser against "
              "overfitting. Their study is a useful illustration of the practical nuance "
              "required when standard architectures are transplanted into medical "
              "imaging.")
    para(doc,
         "Architectural work specific to the domain has also appeared. Erten and "
              "colleagues [12] proposed ConcatNeXt, which fuses multi-scale "
              "representations from parallel convolutional branches so that fine "
              "cytoplasmic texture and coarse nuclear shape are captured simultaneously, "
              "surpassing established backbones including ResNet [25] and InceptionV3 "
              "across several benchmarks. The design rationale is directly cytological: "
              "the features that distinguish leukocyte subtypes exist at different "
              "spatial scales, and a single-scale feature stream must trade one against "
              "the other. Ghaderzadeh and colleagues [13] fused features from DarkNet-53 "
              "and ShuffleNet, selected discriminative components by principal component "
              "analysis, and passed them to a support vector machine, reporting 100% "
              "accuracy on the ALL-IDB benchmark. A perfect score on a small two-class "
              "benchmark raises questions of dataset difficulty and diversity at least as "
              "much as it demonstrates method strength, and this pattern recurs "
              "throughout the literature.")
    para(doc,
         "Subtype-specific work has attracted intense attention, particularly for "
              "acute lymphoblastic leukaemia given its paediatric prevalence. Aria and "
              "colleagues [14] evaluated ten architectures in a two-channel network for "
              "simultaneous detection and subtype classification, with DenseNet201 "
              "achieving 99.85% accuracy and 99.52% sensitivity on 3,562 images from 89 "
              "patients. Rahman and colleagues [15] demonstrated with a lightweight "
              "EfficientNet-B3 using depthwise separable convolutions that heavily "
              "parameterised networks are not always necessary, achieving competitive "
              "performance at reduced computational cost, which matters for deployment in "
              "low-resource settings where the diagnostic need is often greatest. Related "
              "contributions include squeeze-and-excitation frameworks for leukaemia "
              "detection [29], dedicated feature extractors for acute leukaemia diagnosis "
              "from smear images [30], and attention-coupled ResNet and DenseNet hybrids "
              "[28]. Lewicki and colleagues [31] extended the approach to acute "
              "promyelocytic leukaemia in bone marrow smears, a clinically urgent subtype "
              "where rapid identification changes management, and Wang and colleagues "
              "[32] reported evaluation of a diagnosis system inside a real clinical "
              "workflow instead of on a curated benchmark, which remains uncommon.")
    para(doc,
         "Multi-level classification distinguishing both lineage groups and specific "
              "types has been explored architecturally. Acevedo, Merino and colleagues "
              "[16] proposed a two-stage scheme employing region detection to separate "
              "mononuclear from polymorphonuclear types at the first level, then parallel "
              "MobileNet subnetworks for fine-grained classification at the second, "
              "achieving approximately 98.4% across evaluation metrics under Monte Carlo "
              "cross-validation. This sequential coarse-to-fine architecture is an "
              "implicit form of hierarchical inference and is the closest antecedent to "
              "the present work. Its limitation, developed in Section 2.1.6, is that the "
              "two stages are trained independently.")

    h3(doc, "2.1.3  Dataset Benchmarks and Their Limitations")
    para(doc,
         "Large annotated datasets have been the critical enabling factor for all of "
              "the above. In addition to Raabin-WBC [5] and the segmentation ground truth "
              "of Kouzehkanan and colleagues [6], earlier work by Acevedo and colleagues "
              "[26] established peripheral blood benchmarks in routine use. The MLL23 "
              "collection [1] represents the current state of the art, comprising 41,621 "
              "expertly annotated single-cell images across 18 specific classes, and is "
              "the dataset employed here. Its class vocabulary is markedly finer than "
              "most alternatives, resolving the myeloid maturation sequence into distinct "
              "stages instead of collapsing them into one immature category.")
    para(doc,
         "A fundamental limitation of most existing datasets, and of models trained "
              "on them, is an unnatural class distribution. Raabin-WBC, BCCD and LISC are "
              "composed predominantly of mature healthy leukocytes in roughly comparable "
              "proportions, because curation for balance makes benchmarking convenient. "
              "MLL23 retains the natural long-tailed distribution of physiological blood, "
              "with rare pathological types outnumbered in proportions exceeding 260 to "
              "one.")
    para(doc,
         "The consequence is that reported accuracies on balanced benchmarks are not "
              "comparable to performance on clinically representative data, and the "
              "difference is not one of degree but of which failure modes are possible at "
              "all. On a balanced four-class problem, no class can be abandoned without a "
              "visible accuracy penalty. On a corpus where one class holds 0.08% of the "
              "images, abandonment is nearly free, and a model may achieve an excellent "
              "headline score while never once predicting the class that would matter "
              "most clinically. Benchmark curation does not simply make the problem "
              "easier; it removes the specific phenomenon that makes deployment "
              "difficult, which is why this study treats the imbalance as its subject, "
              "not as an obstacle.")

    h3(doc, "2.1.4  Morphological Ambiguity and the Long-Tailed Distribution")
    para(doc,
         "Two challenges dominate the domain. The first is extreme morphological "
              "similarity between biologically distinct types. A myeloblast and a "
              "promyelocyte share nuclear features and cytoplasmic coloration yet "
              "represent different diagnostic entities; large lymphocytes and monocytes "
              "are confused even by experienced staff. The difficulty is intrinsic, not "
              "technological, because haematopoiesis is a continuous developmental "
              "process onto which discrete labels have been imposed. Cells genuinely "
              "exist between the named stages, and the boundary is a convention agreed "
              "among cytologists, not a natural break.")
    para(doc,
         "This inter-class overlap is compounded by within-class variability arising "
              "from donor status, fixation and staining duration, slide thickness, "
              "magnification, exposure and the position of the cell on the slide, where "
              "folding or smearing artefacts distort morphology. Modaghegh and colleagues "
              "[3] catalogued these factors in a systematic review and found that most "
              "classification models are evaluated on curated benchmarks with "
              "near-uniform staining, biasing reported accuracy upward relative to "
              "multi-institution deployment. Rehman and colleagues [18] reached the same "
              "conclusion in an independent review, identifying generalisability across "
              "institutions with differing equipment and protocols as a critical unsolved "
              "problem in computational haematology.")
    para(doc,
         "The second challenge is the severe imbalance of cell type frequencies. In "
              "healthy peripheral blood, neutrophils comprise 55 to 70% of leukocytes, "
              "lymphocytes 20 to 40%, monocytes 2 to 8%, eosinophils 1 to 4%, and "
              "basophils under 1%. Pathological precursors such as blasts and immature "
              "granulocytes appear in proportions more dramatically skewed still, and it "
              "is their presence, not their abundance, that is diagnostically decisive.")
    para(doc,
         "This distribution creates a direct problem for classifiers trained with "
              "unweighted cross-entropy. A classifier minimising global expected loss "
              "achieves its lowest value by learning majority decision boundaries "
              "accurately at the expense of minority features, because the gradient "
              "contributed by a class is proportional to its representation in the batch. "
              "The result is inflated aggregate metrics alongside near-zero sensitivity "
              "for the rarest and most clinically relevant cells. Mushtaq and colleagues "
              "[8] identified this bias as a primary failure mode of existing clinical "
              "tools and noted that it is almost never surfaced by the accuracy metrics "
              "conventionally reported in computer vision studies.")
    para(doc,
         "The clinical implication is asymmetric in a way that symmetric loss "
              "functions do not encode. Failing to detect a leukaemic blast is not "
              "equivalent to an ordinary misclassification; it is a false negative that "
              "may delay a life-saving diagnosis, whereas a false positive prompts a "
              "review that will correct it. Nor are errors equivalent to one another: "
              "confusing two adjacent maturation stages is a mild error that human review "
              "would likely tolerate, while assigning a lymphoid cell to the myeloid "
              "lineage is a categorically different failure. This asymmetry motivates "
              "both the class-balanced objective and the hierarchical error analysis "
              "adopted in this study, and its absence from the reviewed literature is one "
              "of the gaps identified in Section 2.2.")

    h3(doc, "2.1.5  Handling Class Imbalance in Deep Learning")
    para(doc,
         "Two methodological families address imbalance: data-level approaches that "
              "rectify the distribution before training, and algorithm-level approaches "
              "that modify the objective while leaving the data intact.")
    para(doc,
         "Data-level methods include oversampling by SMOTE and ADASYN, undersampling "
              "of majority classes, and generative synthesis of minority images. Abbas "
              "and colleagues [21] incorporated generative augmentation in a hybrid "
              "Inception-XGBoost model reporting 100% accuracy on ALL-IDB, and "
              "Ghaderzadeh and colleagues [13] used generative networks to address "
              "scarcity in a multi-class pipeline. These approaches carry significant "
              "limitations in medical imaging. Interpolation between minority samples in "
              "high-dimensional pixel space frequently produces biologically implausible "
              "artefacts corresponding to no genuine cellular morphology, since the space "
              "of valid cell images is a thin manifold within pixel space and the "
              "straight line between two valid points leaves it immediately. Generative "
              "training is itself unstable and prone to mode collapse on precisely the "
              "small minority sets where it is most needed. And aggressive oversampling "
              "of an extremely rare class causes the network to memorise the few source "
              "images instead of learning features that generalise.")
    para(doc,
         "Khan and colleagues [22] demonstrated this last risk directly, showing "
              "that ADASYN-based resampling followed by feature selection achieved 100% "
              "cross-validation accuracy while failing to preserve that performance on "
              "held-out splits from independent acquisition sites. The result is a "
              "caution against the evaluation protocol as much as against the resampling: "
              "cross-validation performed after resampling leaks synthetic variants of "
              "the same source images across folds, and reports a number that cannot be "
              "reproduced on genuinely unseen data.")
    para(doc,
         "Algorithm-level approaches modify the optimisation objective instead, "
              "leaving the empirical distribution untouched. The foundational mechanism "
              "is cost-sensitive learning, reweighting the per-class loss inversely to "
              "frequency. Naive inverse-frequency weighting is extremely sensitive to "
              "very small class sizes, however, producing enormous weights that "
              "destabilise training and cause the model to overfit a handful of examples. "
              "Cui and colleagues [7] identified this flaw and proposed a class-balanced "
              "loss based on the effective number of samples, reasoning that as a class "
              "grows each additional sample provides diminishing marginal information "
              "because samples overlap in feature space. The resulting reweighting curve "
              "is smooth, bounded and numerically stable, and is directly applicable "
              "where class frequencies span three orders of magnitude as they do here.")
    para(doc,
         "Lin and colleagues [19] introduced focal loss, which attacks imbalance "
              "through difficulty instead of frequency. Easy, confidently correct "
              "examples, which are disproportionately drawn from majority classes, are "
              "down-weighted by a modulating factor so that gradient concentrates on hard "
              "examples. The two mechanisms complement each other; they are not "
              "alternatives: class weighting addresses how many samples a class has, "
              "focal modulation addresses how hard an individual sample is, and the "
              "present study combines them for that reason. Chen and colleagues [20] "
              "proposed the progressive class-centre triplet loss, operating on the "
              "geometry of the learned feature space instead of on softmax probabilities, "
              "repelling class clusters while compacting intra-class distributions so "
              "that minority clusters are not subsumed by the gravitational mass of "
              "majority centroids. This geometric perspective is a conceptual advance "
              "over scalar reweighting, though it introduces additional hyperparameters.")
    para(doc,
         "Stronger long-tail treatments have emerged outside the medical domain and "
              "have not, to the author's knowledge, been evaluated on this task. Kang and "
              "colleagues [36] showed that representation learning and classifier "
              "learning have different optimal sampling regimes, and that learning the "
              "representation under instance-balanced sampling before retraining only the "
              "classifier under class-balanced sampling consistently outperforms "
              "end-to-end reweighting. Menon and colleagues [37] derived logit adjustment "
              "from the statistical consistency of the balanced error, providing a "
              "principled correction applied either during training or post hoc. Both are "
              "computationally inexpensive and both represent plausible improvements on "
              "the treatment adopted here, a point returned to in Chapter 5.")
    para(doc,
         "A limitation is shared by all of these methods. They are designed and "
              "validated for flat, one-dimensional label spaces. None has been formulated "
              "to operate across a multi-level taxonomy. When a model classifies at two "
              "simultaneously optimised levels, it is an open question whether the "
              "imbalance-aware loss should be applied at the coarse level, the fine "
              "level, or through a weighted combination, and whether the effective number "
              "of samples should be computed within or across parent classes. This "
              "dissertation investigates that question directly by applying "
              "class-balanced weighting at both levels, with the lineage-level counts "
              "obtained by summation over each lineage's members.")

    h3(doc, "2.1.6  Hierarchical Classification in Medical Imaging")
    para(doc,
         "Hierarchical frameworks represent taxonomic relationships between "
              "categories explicitly, typically as a tree or directed acyclic graph, "
              "instead of treating every class as equally far from every other. Inference "
              "proceeds coarse to fine, mirroring the reasoning of an expert "
              "haematologist who identifies lineage before distinguishing subtype.")
    para(doc,
         "The theoretical advantages are twofold. Sharing feature representations "
              "between related classes at the coarse level is equivalent to increasing "
              "the effective sample size for the parent decision boundary, which is "
              "particularly valuable for rare classes sharing a lineage with abundant "
              "ones. A class with thirty images cannot support a reliable decision "
              "boundary on its own, but the lineage to which it belongs may have "
              "thousands, and features useful for the lineage are likely useful for its "
              "members. Second, hierarchical architectures alter the cost structure of "
              "misclassification. A flat model penalises confusing a myeloblast with a "
              "lymphocyte identically to confusing it with a promyelocyte, despite the "
              "vastly greater clinical severity of the former, whereas a hierarchical "
              "model with a coarse-level term has an explicit incentive to keep errors "
              "within the correct subtree.")
    para(doc,
         "Empirical support exists but is thin in this domain. Acevedo and "
              "colleagues [16] provided the clearest instantiation for white blood cells, "
              "with mononuclear against polymorphonuclear discrimination at the first "
              "stage followed by within-group classification at the second. Their "
              "reported 98.4% underlines the practical viability of the approach. The "
              "limitation is architectural: the two stages operate independently and are "
              "trained with disjoint loss functions, so the gradient signal from "
              "fine-grained errors does not propagate to inform the coarse discriminator, "
              "and a first-stage error is unrecoverable. The present study differs in "
              "coupling the levels through a single differentiable objective for exactly "
              "this reason.")
    para(doc,
         "An and colleagues [23] demonstrated that enforcing hierarchical structure "
              "during training regularises the model to learn generalised coarse features "
              "before specialising, consistently reducing overfitting on limited minority "
              "samples relative to a flat baseline, in a setting characterised by severe "
              "data scarcity. Kowsari and colleagues [24] reported analogous findings "
              "with the HMIC framework for medical image classification, and observed "
              "that the widest performance gap over flat baselines occurred precisely in "
              "the most severely imbalanced experimental conditions. That observation is "
              "the most direct published support for the hypothesis this dissertation "
              "tests.")
    para(doc,
         "Neither of the latter two studies operates in haematology, however, and "
              "the hierarchies they employ are constructed from visual similarity metrics "
              "or domain-agnostic label graph analysis, not from biological knowledge. A "
              "hierarchy inferred from feature clustering is circular in a way a "
              "biological one is not: it encodes what the model already finds similar, "
              "and so cannot add information the model lacks.")
    para(doc,
         "The opportunity this dissertation exploits is that a scientifically "
              "rigorous hierarchy already exists and is independent of any model. All "
              "blood cells originate from a multipotent haematopoietic stem cell and "
              "bifurcate into myeloid and lymphoid lineages, with immature cells "
              "progressing through well-defined named intermediate stages from blast "
              "through promyelocyte, myelocyte, metamyelocyte and band cell to the mature "
              "form. Wang [17] articulates this structure as a fundamental organising "
              "principle that deep learning models in haematology should not simply "
              "ignore. The present study operationalises it by deriving the taxonomy "
              "directly from the haematopoietic tree, so that the structural prior "
              "supplied to the model is genuine outside knowledge, not a restatement of "
              "what the model already encodes.")

    h3(doc, "2.1.7  Explainable Artificial Intelligence in Medical Diagnostics")
    para(doc,
         "Clinical adoption requires that predictions be auditable. A model "
              "achieving high accuracy by attending to staining artefacts, slide position "
              "or neighbouring erythrocytes will not transfer to a new laboratory, and "
              "cannot reasonably be trusted by the practitioner expected to act on it. "
              "Explainability here is not an ethical supplement but a component of "
              "validation: it is the means by which a shortcut-learning failure is caught "
              "before deployment instead of after.")
    para(doc,
         "Gradient-based saliency methods are the dominant approach. Grad-CAM [9] "
              "produces class-discriminative localisation maps by weighting convolutional "
              "activation maps by the spatially pooled gradient of the target logit and "
              "rectifying the result, requiring no architectural modification and no "
              "retraining, which is why it has been adopted so widely. Islam and "
              "colleagues [10] applied explainable analysis to an optimised convolutional "
              "network for blood cell classification, confirming that saliency settled on "
              "the cell instead of the background and thereby providing evidence that the "
              "reported accuracy was not an artefact of spurious correlation. Al-Dulaimi "
              "and colleagues [28] incorporated attention mechanisms directly into the "
              "architecture, which yields interpretability as a by-product of the forward "
              "pass instead of a separate gradient analysis afterwards.")
    para(doc,
         "Wenderoth and colleagues [27] addressed the related problem of data "
              "limitation through self-supervised pretraining, demonstrating transferable "
              "representations across haematological classification tasks without "
              "requiring expert annotation at scale. Their work is relevant here because "
              "the representations learned without labels are, by construction, not "
              "shaped by the label distribution, and may for that reason be less biased "
              "toward majority classes than supervised features.")
    para(doc,
         "Saliency for hierarchical models remains an open frontier. Where a model "
              "emits predictions at two taxonomic levels, the natural questions are "
              "whether the coarse and fine decisions attend to the same evidence, whether "
              "disagreement between the levels is visible in the saliency, and whether a "
              "cross-lineage error is preceded by a shift in attention away from the "
              "discriminative region. The reviewed literature does not address these "
              "questions, and the present study treats multi-level saliency as a "
              "deliverable, not an illustration, although a full analysis remains beyond "
              "its scope.")

    h2(doc, "2.2  Critical Analysis of Existing Studies")
    para(doc,
         "The studies reviewed above differ along four parameters material to the "
              "present work: the dataset used and whether it preserves clinical "
              "prevalence, the number of classes discriminated, the treatment applied to "
              "imbalance, and whether label structure is exploited. Table 2.1 summarises "
              "the comparison and positions the proposed system against it.")

    table(doc,
          ["Study", "Dataset (classes)", "Method", "Imbalance treatment",
           "Hierarchy", "Reported accuracy"],
          [
              ["Matek et al. [2]", "Bone marrow (21)", "CNN, ResNeXt",
               "None", "No", "Expert-level"],
              ["Alam & Islam [11]", "WBC (4)", "DenseNet-121",
               "None", "No", "98.84%"],
              ["Erten et al. [12]", "Multiple (4-8)", "ConcatNeXt",
               "None", "No", ">98%"],
              ["Ghaderzadeh et al. [13]", "ALL-IDB (2)", "DarkNet + ShuffleNet + SVM",
               "GAN synthesis", "No", "100%"],
              ["Aria et al. [14]", "B-ALL (4)", "DenseNet201, two-channel",
               "None", "No", "99.85%"],
              ["Rahman et al. [15]", "Leukaemia (4)", "EfficientNet-B3",
               "None", "No", "~98%"],
              ["Acevedo et al. [16]", "WBC (8)", "Faster R-CNN + MobileNet",
               "None", "Two-stage, disjoint", "98.4%"],
              ["Al-Dulaimi et al. [28]", "WBC (5)", "ResNet + DenseNet + SCAM",
               "None", "No", ">98%"],
              ["Abbas et al. [21]", "ALL-IDB (2)", "Inception-v3 + XGBoost",
               "GAN synthesis", "No", "100%"],
              ["Khan et al. [22]", "Microarray (2)", "Hybrid logistic vector trees",
               "ADASYN", "No", "100% (CV only)"],
              ["An et al. [23]", "Medical (varies)", "Hierarchical transfer learning",
               "None", "Yes, non-biological", "Improved vs flat"],
              ["Kowsari et al. [24]", "Histology (varies)", "HMIC",
               "None", "Yes, non-biological", "Improved vs flat"],
              ["Wenderoth et al. [27]", "Haematology (varies)", "Self-supervised pretraining",
               "None", "No", "Competitive"],
              ["Cui et al. [7]", "Long-tailed vision", "Class-balanced loss",
               "Effective number", "No", "Method paper"],
              ["Lin et al. [19]", "Object detection", "Focal loss",
               "Difficulty reweighting", "No", "Method paper"],
              ["Chen et al. [20]", "Medical (varies)", "PCCT triplet loss",
               "Feature-space margins", "No", "Method paper"],
              ["This study", "MLL23 (18), IR 261:1",
               "Shared backbone, two heads",
               "Class-balanced + focal", "Yes, biological",
               "See Chapter 4"],
          ],
          cap="Table 2.1  Critical analysis and summary of the existing studies",
          widths=[1.15, 1.05, 1.25, 1.0, 0.85, 0.85], font_size=8)

    para(doc,
         "Four observations follow from the table. First, reported accuracies "
              "cluster tightly above 98%, but almost all are obtained on datasets with "
              "few classes and approximately balanced distributions. The apparent "
              "maturity of the field is partly an artefact of benchmark selection: an "
              "accuracy of 98% on a four-class balanced problem is not evidence about "
              "eighteen-class performance under a 261 to one imbalance, and the two "
              "numbers are not comparable even in principle. Two studies report 100% "
              "accuracy, which is more plausibly read as evidence that the benchmark is "
              "saturated than that the problem is solved.")
    para(doc,
         "Second, where imbalance is addressed at all, it is addressed at the data "
              "level by resampling or generative synthesis, despite the documented "
              "tendency of these methods to overfit the rare classes they are meant to "
              "protect [22]. The algorithm-level methods that avoid this failure mode "
              "[7], [19], [20] originate outside the haematology literature and have not "
              "been systematically applied within it, which represents a straightforward "
              "transfer opportunity.")
    para(doc,
         "Third, and most importantly, the studies that exploit label structure and "
              "the studies that address imbalance are largely disjoint sets. Acevedo and "
              "colleagues [16] use a hierarchy but train the levels independently and do "
              "not treat imbalance; An and colleagues [23] and Kowsari and colleagues "
              "[24] use hierarchies but non-biological ones and outside haematology; Cui "
              "and colleagues [7] and Lin and colleagues [19] treat imbalance but assume "
              "a flat label space. No reviewed study occupies the intersection, which is "
              "where this dissertation is positioned.")
    para(doc,
         "Fourth, a methodological observation concerns evaluation, not method. The "
              "majority of the reviewed studies report accuracy as the headline metric, "
              "with macro-averaged measures given secondarily or not at all. On a "
              "balanced benchmark this is unobjectionable, since accuracy and macro F1 "
              "nearly coincide. On a clinically representative corpus it conceals the "
              "very behaviour of interest, since a model may abandon the rarest class "
              "entirely at negligible cost to accuracy. Chapter 4 demonstrates this "
              "concretely within a single controlled experiment. Furthermore, none of the "
              "reviewed studies decomposes errors by taxonomic distance, so the "
              "clinically important distinction between a mild within-lineage confusion "
              "and a severe cross-lineage one is not measured anywhere in this "
              "literature, and any effect on that distinction in existing systems would "
              "have been invisible to their authors.")
    para(doc,
         "Five conclusions follow from the review. Deep convolutional and "
              "transformer models can recognise blood cells at expert level. Generalising "
              "across different microscopes and laboratories remains a known weakness "
              "that nobody has solved. Class imbalance is widely acknowledged, but the "
              "usual remedies carry a documented risk of overfitting the very classes "
              "they are meant to protect. Hierarchical classification helps in medical "
              "imaging, yet no one has applied it to peripheral blood using a taxonomy "
              "taken from biology. And evaluation practice across the field fits the "
              "benchmarks better than it fits the clinical problem those benchmarks stand "
              "in for.")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run("Research gap.  ")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(12)
    r2 = p.add_run(
        "No existing methodology simultaneously combines a biologically faithful "
             "haematopoiesis-derived hierarchy, an algorithm-level class-balanced "
             "objective operating across multiple taxonomic levels, and saliency "
             "analysis, evaluated on a large peripheral blood single-cell dataset that "
             "preserves genuine clinical imbalance.")
    r2.font.name = FONT
    r2.font.size = Pt(12)

    para(doc,
         "This gap defines the contribution of the present study. The following "
              "chapter describes the methodology adopted to address it, and Chapter 4 "
              "reports the resulting evidence, including the respects in which the "
              "hypothesis was not supported.")
    pagebreak(doc)
