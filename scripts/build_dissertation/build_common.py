"""Shared helpers for building the filled dissertation .docx."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parents[2]
ART = REPO / "artifacts"

FONT = "Times New Roman"


def new_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    for name, size in (("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 13)):
        s = doc.styles[name]
        s.font.name = FONT
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.5
    return doc


def para(doc, text="", *, align=None, bold=False, italic=False, size=None,
         space_after=None, style=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.name = FONT
        if size:
            r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(10)
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    return p


def figure(doc, filename, cap, width_in=5.8):
    path = ART / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    caption(doc, cap)


def table(doc, headers, rows, cap=None, widths=None, font_size=9):
    if cap:
        caption(doc, cap)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(htxt)
        r.bold = True
        r.font.name = FONT
        r.font.size = Pt(font_size)
        hdr[i].paragraphs[0].paragraph_format.line_spacing = 1.0
        hdr[i].paragraphs[0].paragraph_format.space_after = Pt(2)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.name = FONT
            r.font.size = Pt(font_size)
            cells[i].paragraphs[0].paragraph_format.line_spacing = 1.0
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for it in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(it)
        r.font.name = FONT
        r.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)


def pagebreak(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_page_numbers(doc):
    """Page number field in the footer of every section."""
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        for instr, kind in (("begin", "w:fldCharType"), (None, None), ("end", "w:fldCharType")):
            pass
        fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
        run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
        run.font.name = FONT
        run.font.size = Pt(11)
