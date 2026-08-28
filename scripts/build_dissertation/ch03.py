from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from build_common import *


def eq(doc, text, number):
    """Centred display equation with a right-aligned equation number."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.name = FONT
    r.font.size = Pt(12)
    r2 = p.add_run("\t\t(" + number + ")")
    r2.font.name = FONT
    r2.font.size = Pt(12)
    return p


def chapter3(doc):
    h1(doc, "Chapter 3  Methodology")
    para(doc,
         "This chapter describes the architecture of the proposed system, the "
              "collection and preprocessing of data, the model and its training "
              "objective, the metrics used to evaluate it, and the experimental design "
              "under which the comparisons are made.")
    para(doc,
         "The design is comparative and ablative. A single implementation trains "
              "every variant, and the variants differ only in the values of a "
              "configuration object, never in code path. This matters because the effects "
              "under investigation are small. Any incidental divergence between a "
              "baseline and a proposed model would be large enough to manufacture or "
              "destroy the result, so holding everything constant except the factor under "
              "test is a precondition for the comparison being interpretable.")
    figure(doc, "fig_architecture.png",
           "Figure 3.1  Architecture of the proposed system", width_in=6.0)
    para(doc,
         "Figure 3.1 shows the pipeline end to end. Images pass through a "
              "deterministic preprocessing stage into a pretrained backbone shared by two "
              "classification heads. A single composite loss backpropagates through both "
              "heads into the backbone. The trained model emits a combined lineage and "
              "cell-type prediction, and is evaluated with hierarchical error analysis "
              "and gradient-based saliency.")

    # ------------------------------------------------ 3.1
    h2(doc, "3.1  Data Collection and Preprocessing")

    h3(doc, "3.1.1  Dataset and label hierarchy")
    para(doc,
         "The study uses the MLL23 peripheral blood single-cell dataset [1], "
              "distributed through Zenodo under DOI 10.5281/zenodo.14277609. It contains "
              "41,621 Pappenheim-stained single nucleated-cell images at 288 by 288 "
              "pixels, each annotated by expert cytomorphologists into one of 18 classes. "
              "Archives are MD5-verified on download and realised per-class counts "
              "asserted against the data descriptor, since a truncated download or "
              "mis-mapped class directory would corrupt every downstream result while "
              "leaving a pipeline that still runs.")
    para(doc,
         "Each image carries a joint target consisting of a lineage label y1 in "
              "{0,1,2} and a fine cell-type label y2 in {0,...,17}. Because every fine "
              "class belongs to exactly one lineage, the levels are related by a fixed, "
              "non-learned surjection.")
    eq(doc, "y1 = π(y2),   π : {0,...,17} → {0,1,2}", "3.1")
    para(doc,
         "Equation (3.1) is the structural prior the study exploits. It is known "
              "biology, not something the model must discover. Fine-class indices are "
              "assigned lineage by lineage, and within the myeloid branch they follow the "
              "maturation continuum from myeloblast through promyelocyte, myelocyte, "
              "metamyelocyte and band neutrophil to segmented neutrophil. This ordering "
              "is preserved in the class indices and in every confusion-matrix axis, so "
              "biologically adjacent stages appear adjacent to the diagonal. "
              "Adjacent-stage confusion is the clinically expected error mode, and an "
              "alphabetical or frequency-sorted axis would scatter it across the plot.")
    figure(doc, "fig_hierarchy.png",
           "Figure 3.2  The two-level label hierarchy", width_in=5.6)
    figure(doc, "class_examples.png",
           "Figure 3.3  One representative image per class after preprocessing, with "
                "corpus counts", width_in=6.2)
    para(doc,
         "Figure 3.3 makes the difficulty visible directly: the myeloid maturation "
              "stages differ by gradual changes in nuclear shape, chromatin condensation "
              "and cytoplasmic granularity, and by no single categorical feature.")

    h3(doc, "3.1.2  Class imbalance")
    para(doc,
         "Writing n_c for the number of images in class c, the imbalance ratio is "
              "the quotient of the largest and smallest class counts.")
    eq(doc, "IR = max(n_c) / min(n_c) = 8,606 / 33 ≈ 261", "3.2")
    para(doc,
         "Eight classes fall at or below a threshold of 1,000 images and are "
              "designated minority classes throughout. The threshold sits in a natural "
              "gap in the distribution, the next class above it holding 1,658 images, so "
              "it separates a real cluster instead of cutting one in half. This imbalance "
              "is the subject matter of the study, not an obstacle to be removed, and it "
              "is never resampled away. The training distribution is left intact and the "
              "imbalance is addressed in the objective function. Resampling would change "
              "the question being asked.")
    figure(doc, "class_distribution.png",
           "Figure 3.4  Class distribution on linear and logarithmic axes, coloured by "
                "lineage", width_in=6.2)

    h3(doc, "3.1.3  Deterministic stratified partitioning")
    para(doc,
         "The corpus is partitioned once into training, validation and test sets in "
              "70/15/15 proportion, stratified on the fine label. Stratifying on the fine "
              "label stratifies the lineage automatically by Equation (3.1). The "
              "partition is produced by two successive stratified draws under a fixed "
              "seed and then persisted to disk, so every arm, seed and backbone reads the "
              "same file and the split is identical by construction, not by convention. "
              "Table 3.1 gives the realised counts.")
    para(doc,
         "Three conditions are checked in code, not assumed: that no class is absent "
              "from any partition, that no image path occurs in more than one, and that "
              "realised fractions deviate from requested ones by less than one percentage "
              "point. The first matters because the rarest class holds only 33 images, "
              "leaving 23, 5 and 5 under a 70/15/15 split. An empty validation or test "
              "cell would make its recall undefined and silently distort the macro "
              "average.")

    rows = [
        ["0", "Typical lymphocytes", "Lymphoid", "5,532", "3,872", "830", "830"],
        ["1", "Hairy cells", "Lymphoid", "3,265", "2,285", "490", "490"],
        ["2", "Large granular lymphocytes", "Lymphoid", "1,849", "1,294", "277", "278"],
        ["3", "Atypical lymphocytes (plasma cells)", "Lymphoid", "1,658", "1,161", "249", "248"],
        ["4", "Smudge cells †", "Lymphoid", "988", "692", "148", "148"],
        ["5", "Neoplastic lymphocytes †", "Lymphoid", "180", "126", "27", "27"],
        ["6", "Reactive lymphocytes †", "Lymphoid", "33", "23", "5", "5"],
        ["7", "Myeloblasts", "Myeloid", "8,606", "6,024", "1,291", "1,291"],
        ["8", "Promyelocytes †", "Myeloid", "745", "521", "112", "112"],
        ["9", "Atypical promyelocytes", "Myeloid", "2,033", "1,423", "305", "305"],
        ["10", "Myelocytes †", "Myeloid", "747", "523", "112", "112"],
        ["11", "Metamyelocytes †", "Myeloid", "483", "338", "73", "72"],
        ["12", "Band neutrophils †", "Myeloid", "687", "481", "103", "103"],
        ["13", "Segmented neutrophils", "Myeloid", "7,170", "5,019", "1,075", "1,076"],
        ["14", "Eosinophil granulocytes", "Myeloid", "2,448", "1,714", "367", "367"],
        ["15", "Basophil granulocytes †", "Myeloid", "616", "431", "93", "92"],
        ["16", "Monocytes", "Myeloid", "2,510", "1,757", "376", "377"],
        ["17", "Normoblasts", "Erythroid", "2,071", "1,450", "310", "311"],
        ["", "Total", "", "41,621", "29,134", "6,243", "6,244"],
    ]
    table(doc, ["y2", "Cell type", "Lineage", "Total", "Train", "Val", "Test"],
          rows,
          cap="Table 3.1  The 18 MLL23 classes, lineage membership and partition sizes. "
               "Minority classes (n ≤ 1,000) are marked †",
          widths=[0.35, 2.0, 0.8, 0.7, 0.7, 0.55, 0.55], font_size=8)

    h3(doc, "3.1.4  Preprocessing and augmentation")
    para(doc,
         "Evaluation preprocessing is deterministic. An image is bilinearly "
              "resampled from 288 to 224 pixels, scaled to the unit interval, and "
              "standardised channel-wise using ImageNet statistics.")
    eq(doc, "x̂ = (x/255 − μ) / σ,   μ = (0.485, 0.456, 0.406),  σ = (0.229, 0.224, "
         "0.225)", "3.3")
    para(doc,
         "Two choices deserve comment. The input resolution is 224 by 224 for every "
              "backbone including the Vision Transformer, since holding resolution "
              "constant is what keeps the backbone comparison controlled. And ImageNet "
              "statistics are retained even though the corpus channel means are brighter "
              "in every channel, so the standardised data centres away from zero. This is "
              "intentional: pretrained backbones expect inputs standardised the way they "
              "were trained, and dataset-specific statistics would trade away part of the "
              "transfer benefit.")
    para(doc,
         "Training applies online augmentation before this deterministic pipeline. "
              "It is class-conditional, so the eight rarest types are presented in "
              "greater apparent variety without altering the sampling distribution: a "
              "strength multiplier of 2.0 applies to rotation range and jitter for "
              "minority classes and 1.0 otherwise, capped as in Table 3.2.")
    table(doc, ["Transform", "Majority classes", "Minority classes"], [
        ["Horizontal flip", "p = 0.5", "p = 0.5"],
        ["Vertical flip", "p = 0.5", "p = 0.5"],
        ["Random rotation", "±90°", "±180°"],
        ["Brightness / contrast jitter", "0.1", "0.2"],
        ["RandAugment", "2 ops, magnitude 9", "3 ops, magnitude 9"],
    ], cap="Table 3.2  Class-conditional augmentation parameters",
        widths=[2.2, 1.7, 1.7], font_size=10)
    para(doc,
         "Rotation is label-preserving here in a way it is not for natural images, "
              "since a blood cell has no canonical orientation. It is performed at native "
              "resolution before the downsample, so the image is resampled once instead "
              "of twice, and the corners are filled with white to match the smear "
              "background. Figure 3.5 exposes a caveat: that fill applies to the explicit "
              "rotation only, whereas the automated policies perform their own geometric "
              "operations defaulting to black and can invert the stain palette outright. "
              "This is a recognised property of such policies, not a defect, since their "
              "regularising value derives from the severity of the perturbation, but it "
              "qualifies the justification above.")
    figure(doc, "augmentation_policies.png",
           "Figure 3.5  The three augmentation policies applied to one image, four "
                "draws each", width_in=5.4)

    h3(doc, "3.1.5  Stain normalisation")
    para(doc,
         "Slide preparation and imaging introduce colour variation unrelated to cell "
              "identity. One experimental arm evaluates stain normalisation as an "
              "additional preprocessing stage using the method of Reinhard [33], in which "
              "the image is converted to the decorrelated Lab colour space and each "
              "channel matched to a reference by mean and standard deviation.")
    eq(doc, "L′ = (L − mean(L)) × (σ_ref / σ) + mean(L_ref)", "3.4")
    para(doc,
         "Reinhard is preferred to Macenko [34] on cost grounds, being roughly "
              "twenty times cheaper; per-image stain-basis estimation would starve the "
              "accelerator and make that arm slower for reasons unrelated to the science. "
              "Critically, the reference statistics are fitted on the training split "
              "alone, since fitting on the full corpus would leak validation and test "
              "colour statistics into training.")
    figure(doc, "stain_comparison.png",
           "Figure 3.6  Stain normalisation, one image per lineage: original, Macenko, "
                "Reinhard", width_in=4.4)

    # ------------------------------------------------ 3.2
    h2(doc, "3.2  ML/AI Model Development")

    h3(doc, "3.2.1  Shared backbone with two heads")
    para(doc,
         "The model is a pretrained backbone producing a pooled D-dimensional "
              "representation, followed by two linear heads emitting lineage logits u and "
              "fine logits v respectively.")
    eq(doc, "z = Dropout(fθ(x̂));   u = W₁z + b₁ ∈ ℝ³;   v = W₂z + b₂ ∈ ℝ¹⁸", "3.5")
    para(doc,
         "The heads are trained jointly on the shared trunk, so the coarse decision "
              "regularises the fine one and rare types borrow representational strength "
              "from their lineage neighbours. A single class implements every arm: in "
              "flat mode the lineage head is never constructed, so it contributes neither "
              "parameters nor weight decay, an unused head having made the claim of an "
              "identical backbone quietly false. At inference the flat model's lineage "
              "prediction is derived from its fine prediction through Equation (3.1), "
              "which makes the hierarchical error decomposition computable for the "
              "baseline and hence comparable across arms.")
    para(doc,
         "An obvious alternative is a hard gate that predicts lineage first and "
              "routes the sample to one of three per-lineage subclassifiers. It was "
              "rejected on evidence. Exploratory analysis on frozen ImageNet features "
              "measured a lineage silhouette coefficient of approximately zero, so the "
              "lineages form no separable global clusters, while k-nearest-neighbour "
              "lineage agreement was 0.78 against a chance baseline of 0.45. The "
              "structure is real but local. A hard gate assumes global separability, "
              "which the data rejects, and every misfire would produce an unrecoverable "
              "cross-lineage error. The soft coupling adopted instead exploits the local "
              "structure without asserting the global one.")
    figure(doc, "umap_lineage.png",
           "Figure 3.7  UMAP projection of frozen backbone features, coloured by "
                "lineage", width_in=4.6)
    para(doc,
         "Figure 3.7 is a two-dimensional projection computed on frozen features "
              "before fine-tuning, so it motivates the architectural choice without "
              "bounding what the trained model can represent.")
    table(doc, ["Key", "Architecture", "Family"], [
        ["mobilenet", "MobileNetV3-Large", "Lightweight CNN"],
        ["resnet", "ResNet-50", "Residual CNN"],
        ["convnext", "ConvNeXt-Tiny", "Modernised CNN"],
        ["vit", "ViT-Small/16", "Vision Transformer"],
    ], cap="Table 3.3  Candidate backbone architectures, all fine-tuned end to end at "
         "224 by 224",
        widths=[1.3, 2.2, 2.0], font_size=10)

    h3(doc, "3.2.2  The imbalance-robust hierarchical objective")
    para(doc,
         "The objective has three components. Two address imbalance and the third "
              "couples the label levels. Following Cui and colleagues [7], samples of the "
              "same class overlap in feature space, so the n-th sample contributes less "
              "new information than the first. The effective number of samples of a class "
              "is used to derive its weight, which is then normalised to unit mean.")
    eq(doc, "w_c = (1 − β) / (1 − β^n_c),   β = 0.9999", "3.6")
    para(doc,
         "The normalisation to unit mean is essential to the ablation design: "
              "without it the weighted arm would train at a different effective learning "
              "rate, confounding reweighting with step size. The same construction is "
              "applied at the lineage level, which is itself imbalanced with myeloid "
              "accounting for roughly 63% of the corpus.")
    para(doc,
         "Class weighting addresses how many samples a class has, not how hard an "
              "individual sample is. The focal loss of Lin and colleagues [19] supplies "
              "the second factor, down-weighting confidently correct examples so the "
              "gradient is not dominated by the 8,606 myeloblasts.")
    eq(doc, "L_focal(v, y) = (1 − p_y)^γ · CE(v, y; w),   γ = 2", "3.7")
    para(doc,
         "The modulating factor is computed from the unweighted posterior, since it "
              "should reflect how hard the sample is, not how rare its class is; the "
              "class weight already carries rarity and folding it into both factors would "
              "double-count it.")
    para(doc,
         "The term that makes the objective hierarchical instead of merely "
              "multi-task couples the two heads. Let M be the one-hot matrix of the map "
              "in Equation (3.1). Right-multiplying a fine posterior by M marginalises it "
              "into lineage space, since each fine class belongs to exactly one lineage. "
              "The consistency penalty is the Kullback-Leibler divergence between that "
              "marginalised belief and the lineage head's own posterior.")
    eq(doc, "L_cons = KL( softmax(v)ᵀ M  ‖  softmax(u) )", "3.8")
    para(doc,
         "Without Equation (3.8) the heads are free to contradict each other, and "
              "the claim that the coarse decision regularises the fine one would have no "
              "mechanism behind it. Gradients flow into both heads, pulling them toward "
              "agreement instead of forcing one to chase the other. The three components "
              "combine into the total objective.")
    eq(doc, "L = λ_lin · L_focal(u, y1) + λ_fine · L_focal(v, y2) + λ_cons · L_cons", "3.9")
    para(doc,
         "with weights 0.3, 1.0 and 0.1 respectively. The fine term is held at unity "
              "because it is the task being evaluated. The consistency weight is "
              "deliberately small: the empirical support is for local lineage structure, "
              "so consistency is applied as a nudge, not a constraint, and forcing it "
              "hard would impose a global structure the data does not exhibit. The two "
              "ablations are field values on one object instead of separate loss classes, "
              "which is what keeps the arms honest.")
    table(doc, ["Switch", "Effect", "Purpose"], [
        ["use_hierarchy = False", "λ_lin = λ_cons = 0",
         "Fine-level loss only; removes the hierarchy"],
        ["use_imbalance = False", "w ≡ 1, γ = 0, no smoothing",
         "Plain cross-entropy; removes the imbalance treatment"],
    ], cap="Table 3.4  Ablation switches implemented on the loss object",
        widths=[1.6, 1.7, 2.4], font_size=10)
    para(doc,
         "Batch mixing is applied as a further regulariser. Because the dataset "
              "carries two label levels, the same mixing coefficient applies to both "
              "targets, so the coarse and fine losses see the same blend; independent "
              "mixes per level would let the objective disagree with itself about how "
              "much of each example is targets, which keeps the class-balanced weighting "
              "well defined.")

    # ------------------------------------------------ 3.3
    h2(doc, "3.3  Evaluation of the Proposed System")
    para(doc,
         "Plain accuracy is explicitly rejected as a headline number. At an "
              "imbalance ratio of 261 it cannot see the behaviour under study, since a "
              "classifier abandoning the rarest class entirely forfeits under 0.1% of "
              "accuracy. The headline metrics are macro F1 and balanced accuracy, which "
              "weight every class equally regardless of frequency.")
    eq(doc, "macro-F1 = (1/C) Σ 2 P_c R_c / (P_c + R_c);   balanced accuracy = (1/C) Σ "
         "R_c", "3.10")
    para(doc,
         "Accuracy, macro precision, macro recall and the full confusion matrix are "
              "also reported. Where a rare class receives no predictions its score is "
              "recorded as zero, not left undefined, since a missing value would "
              "propagate into the macro average. A minority-class view reports the same "
              "quantities restricted to samples whose true class is one of the eight "
              "minority classes; it is recall-oriented by construction and is labelled as "
              "such instead of being presented as an overall score.")
    para(doc,
         "Beyond these, every prediction is assigned to exactly one of three "
              "outcomes, using Equation (3.1) to derive lineages from fine labels: "
              "correct; a within-lineage error, where the fine class is wrong but the "
              "lineage right; or a cross-lineage error. The diagnostic of interest is the "
              "share of errors that remained inside the correct lineage.")
    eq(doc, "within-error share = #{within-lineage errors} / #{errors}", "3.11")
    para(doc,
         "Equation (3.11) is the direct test of this study's thesis. A lineage-aware "
              "model should shift errors from cross-lineage to within-lineage even where "
              "total error is unchanged, and the two are not clinically equivalent. No "
              "scalar accuracy measure expresses this distinction, and no reviewed study "
              "reports it.")
    para(doc,
         "Saliency is assessed with Grad-CAM [9], which weights the activation maps "
              "of a chosen layer by the spatially pooled gradient of the target logit and "
              "rectifies the result. Applying this to a Vision Transformer requires care, "
              "since the classifier reads only the class token and patch gradients at the "
              "final block are exactly zero. Saliency is computed instead at the input to "
              "the final block's attention, where patch tokens still route into the class "
              "token.")

    # ------------------------------------------------ 3.4
    h2(doc, "3.4  Experimental Design and Statistical Protocol")
    para(doc,
         "Five experiments are required under identical partitions and backbones. "
              "They reduce to four distinct configurations, because the hierarchy "
              "ablation applied to the proposed model is the flat baseline: removing the "
              "lineage head and its loss terms leaves a single-head classifier. That "
              "configuration is trained once and reported twice, since training it under "
              "two names would consume a fifth of the compute budget to produce two draws "
              "from the same distribution and invite the reader to mistake them for "
              "independent evidence.")
    table(doc, ["Arm", "Mode", "Hierarchy", "Imbalance", "Stain", "Reported as"], [
        ["flat_baseline", "flat", "—", "yes", "—",
         "1. Flat baseline; 4. Hierarchy ablation"],
        ["hierarchical", "hier", "yes", "yes", "—",
         "2. Full hierarchical model (proposed)"],
        ["no_imbalance", "hier", "yes", "—", "—",
         "3. Imbalance ablation"],
        ["stain_norm", "hier", "yes", "yes", "yes",
         "5. Stain-normalised input"],
    ], cap="Table 3.5  The four distinct configurations and the five reported "
         "experiments",
        widths=[1.1, 0.6, 0.75, 0.8, 0.5, 2.0], font_size=9)
    para(doc,
         "The work proceeds in two phases. Phase 1 screens the four backbones of "
              "Table 3.3 under the hierarchical configuration for a short fixed budget at "
              "one seed, selecting the backbone carried into Phase 2 on validation macro "
              "F1. This is a selection run and not the backbone comparison result, since "
              "a short budget need not rank backbones the way a full run would. Phase 2 "
              "trains the four configurations at three seeds each on the selected "
              "backbone, giving twelve runs.")
    para(doc,
         "Because the design uses one fixed partition instead of k folds, paired "
              "samples are obtained by retraining each arm at three seeds on that fixed "
              "partition and pairing arms by seed. This keeps the split identical across "
              "arms, which the ablation design depends on, at the cost of measuring "
              "variation between training runs, not between datasets.")
    eq(doc, "t = mean(d) / (s_d / √n);   Cohen's d_z = mean(d) / s_d", "3.12")
    para(doc,
         "Effect sizes are reported alongside p-values throughout. With three seeds "
              "the test carries two degrees of freedom, which is thin, and reporting the "
              "effect size and the per-seed values prevents the reader from being asked "
              "to lean on the p-value alone.")
    para(doc,
         "Determinism is enforced at three points: the partition is seeded and "
              "persisted, each run seeds Python, NumPy and Torch from its configuration, "
              "and every configuration is written to the results file alongside its "
              "metrics so a recorded run can be reconstructed exactly. Correctness is "
              "enforced by assertion instead of inspection, with the class hierarchy "
              "validating its own invariants at import and the partition checked for "
              "empty cells, duplicates and proportion drift.")
    pagebreak(doc)
