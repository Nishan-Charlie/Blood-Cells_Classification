from docx.shared import Pt, Inches
from build_common import *


def chapter5(doc):
    h1(doc, "Chapter 5  Conclusion and Future Work")

    h2(doc, "5.1  Conclusion")
    para(doc,
         "This study asked whether a lineage-aware hierarchical classifier, combined "
              "with an imbalance-robust training objective, improves fine-grained "
              "classification of single peripheral blood cell images, especially for rare "
              "cell types, relative to a flat transfer-learning classifier on the same "
              "data. Four configurations were trained on the MLL23 corpus of 41,621 "
              "expert-annotated images across 18 classes at an imbalance ratio of 261:1, "
              "under identical partitions, backbones and schedules, at three seeds each.")
    para(doc,
         "The answer is a qualified yes, with the qualification carrying as much "
              "information as the affirmation.")
    para(doc,
         "The proposed hierarchical model outperformed the flat baseline on every "
              "metric measured, achieving a test macro F1 of 0.879 against 0.874, a "
              "minority-class macro F1 of 0.840 against 0.823, a minority recall of 0.799 "
              "against 0.776, and a balanced accuracy of 0.894 against 0.883. Critically, "
              "the improvement was larger on the rare classes than on the corpus as a "
              "whole, which is the pattern the design predicted, not a uniform lift. None "
              "of these aggregate differences reached statistical significance at three "
              "seeds, although the effect sizes were large, with Cohen's d_z above 1.0 "
              "for minority recall and balanced accuracy. On the aggregate metrics the "
              "hierarchical advantage is suggested but not demonstrated.")
    para(doc,
         "One comparison did reach significance, and it was the one identified in "
              "advance as the direct test of the mechanism. The hierarchical model "
              "reduced cross-lineage error, the clinically severe category, from 0.0119 "
              "to 0.0115, consistently across all three seeds, at p = 0.020 with d_z = "
              "-4.04, and correspondingly raised the share of errors remaining within the "
              "correct lineage from 0.726 to 0.733. Total accuracy for the two arms was "
              "identical to three decimal places. The hierarchy does what it was designed "
              "to do, redistributing errors from the severe category toward the mild one, "
              "but the magnitude is small: the absolute reduction corresponds to two or "
              "three images in a test partition of 6,244. This is a reliably detected "
              "small effect, not a clinically decisive one, and it is reported as such.")
    para(doc,
         "The ablation design produced the study's clearest finding, and it concerns "
              "the other component. Removing the class-balanced and focal terms cost "
              "0.076 in minority recall (p = 0.024), 0.032 in balanced accuracy (p = "
              "0.028) and 0.060 in minority macro F1 (p = 0.052), with effect sizes above "
              "two standard deviations throughout. The imbalance objective contributes "
              "roughly an order of magnitude more to rare-class performance than the "
              "hierarchy does. Where credit for the combined system is apportioned, it "
              "belongs predominantly to the loss function, not to the architecture.")
    para(doc,
         "Diagnostic analysis explains why the architectural contribution was "
              "modest, and this is arguably the most useful outcome of the work. Lineage "
              "accuracy saturates at 98.9% and both auxiliary loss terms collapse to "
              "approximately 0.0007 by the end of training. The auxiliary task "
              "regularises an axis the model has already solved, and so supplies almost "
              "no gradient during the period when the hard discriminations, between "
              "myeloid maturation stages and among the rare lymphoid types, are actually "
              "being learned. Lineage supervision cannot help with those distinctions by "
              "construction, since the confusable classes share a lineage label. This is "
              "a design finding, not a tuning deficiency, and it points directly at the "
              "remedy set out below.")
    para(doc,
         "Per-class analysis sharpened the picture considerably, and produced the "
              "most specific confirmation of the mechanism the study obtained. The "
              "hierarchical advantage is not diffuse: of eighteen classes, fifteen differ "
              "from the flat baseline by 0.006 or less in F1, and the aggregate "
              "improvement is carried almost entirely by reactive lymphocytes, which gain "
              "0.070, and neoplastic lymphocytes, which gain 0.050. These are the two "
              "rarest lymphoid types, and they sit in a lineage containing typical "
              "lymphocytes and hairy cells in abundance. The effect appears exactly where "
              "the theory says rare classes should be able to borrow strength from "
              "populous same-lineage neighbours, and is absent where the theory offers no "
              "reason to expect it. Correspondingly, the myeloid maturation stages show "
              "no measurable benefit, the mean difference across the six chain classes "
              "being -0.005 against a seed standard deviation of 0.011. Their confusable "
              "neighbours share their lineage label, so lineage supervision cannot "
              "discriminate among them even in principle. The architecture helps along "
              "one axis of the taxonomy and is inert along another, and describing it as "
              "a general improvement would misrepresent the evidence.")
    para(doc,
         "A fourth finding was negative. Stain normalisation, evaluated as an "
              "additional preprocessing stage, underperformed the proposed model on every "
              "metric, though not significantly. Either colour carries genuine diagnostic "
              "signal in Pappenheim-stained material, or the single-source corpus offers "
              "little inter-site variation for normalisation to correct. The experiment "
              "does not distinguish these, and the result should not be generalised to "
              "multi-site data.")
    para(doc,
         "A fifth observation concerns difficulty, not method, and holds across "
              "every configuration tested. The five hardest classes are reactive "
              "lymphocytes, metamyelocytes, neoplastic lymphocytes, myelocytes and band "
              "neutrophils. Rarity alone does not explain this grouping: basophil "
              "granulocytes, with 92 test images, reach an F1 of 0.982, while "
              "metamyelocytes with 72 reach 0.677. What separates them is that basophils "
              "are visually distinctive whereas metamyelocytes sit between two adjacent "
              "stages of a continuous developmental process. Difficulty here is driven as "
              "much by morphological adjacency as by sample count. That reframes the "
              "problem in a useful way. The tail is hard partly because it is small, and "
              "more data would help. The maturation continuum is hard for a different "
              "reason: the labels draw sharp lines across a process that has none, and no "
              "quantity of extra images would change that.")
    para(doc,
         "Two methodological contributions accompany the empirical ones. The study "
              "demonstrates, from within a single controlled experiment, that accuracy is "
              "the wrong headline metric for this problem: the arm scoring highest on "
              "accuracy at 0.9602 scored lowest on minority macro F1 at 0.7799. And it "
              "establishes the hierarchical decomposition of errors as a measurement that "
              "scalar metrics cannot replace, since it detected a consistent, significant "
              "effect between two arms whose overall accuracy was indistinguishable. The "
              "study also records the invalidation of an initial experimental matrix in "
              "which measurement noise exceeded the effects under study, and the protocol "
              "changes that reduced that noise by 54%, without which none of the "
              "comparisons reported here would have been resolvable.")

    h2(doc, "5.2  Future Work")
    para(doc,
         "Four directions follow from these findings, ordered by the strength of the "
              "evidence supporting them.")
    para(doc, "")
    p = para(doc, "")
    r = p.add_run("Retarget the auxiliary task. ")
    r.bold = True; r.font.name = FONT; r.font.size = Pt(12)
    r2 = p.add_run(
        "The saturation of the lineage head at 98.9% is the clearest diagnosed "
             "cause of the small hierarchical effect, and the most direct remedy is to "
             "define the auxiliary task on an axis that remains unresolved late in "
             "training. Maturation stage within the myeloid branch is the natural "
             "candidate, since that is where the residual confusion demonstrably "
             "resides. An alternative is to factor inference conditionally, so that the "
             "coarse decision genuinely constrains the fine one, instead of two heads "
             "sitting side by side on a shared trunk and being nudged toward agreement. "
             "Both preserve the biological grounding that motivated the design while "
             "directing the supervisory signal at a harder problem.")
    r2.font.name = FONT; r2.font.size = Pt(12)

    p = para(doc, "")
    r = p.add_run("Adopt stronger long-tail methods. ")
    r.bold = True; r.font.name = FONT; r.font.size = Pt(12)
    r2 = p.add_run(
        "Since the imbalance objective proved load-bearing, improving it is the "
             "highest-value direction. Class-balanced reweighting combined with focal "
             "modulation is a comparatively mild treatment. Decoupled training, in which "
             "the representation is learned under instance-balanced sampling and only "
             "the classifier subsequently retrained under class-balanced sampling [36], "
             "is consistently among the strongest reported long-tail methods and is "
             "inexpensive here because it reuses backbones already trained. Logit "
             "adjustment [37] and margin-based objectives enforcing larger decision "
             "margins for rare classes are drop-in alternatives within the existing loss "
             "implementation. The target is to move minority recall off the 0.80 plateau "
             "observed across every arm in this study.")
    r2.font.name = FONT; r2.font.size = Pt(12)

    p = para(doc, "")
    r = p.add_run("Raise statistical power. ")
    r.bold = True; r.font.name = FONT; r.font.size = Pt(12)
    r2 = p.add_run(
        "Three seeds give two degrees of freedom, and several comparisons in this "
             "study showed large effect sizes while remaining formally unresolved. Five "
             "seeds, combined with the halved measurement noise already achieved, would "
             "convert the paired tests from indicative to informative. A full-length "
             "backbone comparison is also outstanding: the screen conducted here "
             "selected an architecture but was too short to compare them, so the choice "
             "of Vision Transformer rests on weaker evidence than the rest of the "
             "design.")
    r2.font.name = FONT; r2.font.size = Pt(12)

    p = para(doc, "")
    r = p.add_run("Address cross-site generalisation. ")
    r.bold = True; r.font.name = FONT; r.font.size = Pt(12)
    r2 = p.add_run(
        "The most significant limitation of this study is that MLL23 originates "
             "from a single laboratory, so the generalisation problem repeatedly "
             "identified in the literature as the field's central unsolved issue [3], "
             "[18] is untested here. Evaluating the trained model on an independently "
             "acquired corpus such as Raabin-WBC [5] would establish whether the "
             "hierarchical coupling confers robustness to acquisition shift, and would "
             "also resolve the ambiguity in the stain normalisation result, which may "
             "behave quite differently where genuine inter-site colour variation exists. "
             "Self-supervised pretraining on unlabelled smear images [27] is a "
             "complementary route to the same objective.")
    r2.font.name = FONT; r2.font.size = Pt(12)

    para(doc,
         "Beyond these, two smaller refinements were identified during "
              "implementation. The automated augmentation policies fill geometric "
              "transformations with black and can invert the stain palette, producing "
              "images far outside the distribution of any real smear; constraining them "
              "to a white fill consistent with the smear background is a single-parameter "
              "change worth evaluating. And multi-level saliency, comparing where the "
              "coarse and fine heads attend and whether their disagreement is visible, "
              "remains an open question that the present analysis raises without "
              "answering.")
    pagebreak(doc)


def references(doc):
    h1(doc, "References")
    para(doc, "References follow IEEE style.", italic=True)
    refs = [
        "S. Shetab Boushehri, S. Kazeminia, A. Gruber et al., “A large "
             "expert-annotated single-cell peripheral blood dataset for hematological "
             "disease diagnostics,” Scientific Data, vol. 12, 1773, 2025.",

        "C. Matek, S. Krappe, C. Münzenmayer, T. Haferlach and C. Marr, “Highly "
             "accurate differentiation of bone marrow cell morphologies using deep "
             "neural networks on a large image data set,” Blood, vol. 138, no. 20, pp. "
             "1917–1927, 2021.",

        "M. H. S. Modaghegh et al., “Machine learning in detection and "
             "classification of leukemia using smear blood images: a systematic review,” "
             "Scientific Programming, vol. 2021, 9933481, 2021.",

        "C. Matek, S. Schwarz, K. Spiekermann and C. Marr, “Human-level recognition "
             "of blast cells in acute myeloid leukaemia with convolutional neural "
             "networks,” Nature Machine Intelligence, vol. 1, pp. 538–544, 2019.",

        "Z. M. Kouzehkanan, S. Saghari, S. Tavakoli et al., “A large dataset of "
             "white blood cells containing cell locations and types, along with "
             "segmented nuclei and cytoplasm,” Scientific Reports, vol. 12, 1123, 2022.",

        "S. M. Z. Kouzehkanan et al., “New segmentation and feature extraction "
             "algorithm for classification of white blood cells in peripheral smear "
             "images,” Scientific Reports, vol. 11, 19428, 2021.",

        "Y. Cui, M. Jia, T.-Y. Lin, Y. Song and S. Belongie, “Class-balanced loss "
             "based on effective number of samples,” in Proc. IEEE/CVF Conf. Computer "
             "Vision and Pattern Recognition, 2019, pp. 9268–9277.",

        "Z. Mushtaq et al., “A systematic review on recent advancements in deep and "
             "machine learning based detection and classification of acute lymphoblastic "
             "leukemia,” IEEE Access, vol. 10, pp. 90755–90776, 2022.",

        "R. R. Selvaraju, M. Cogswell, A. Das, R. Vedantam, D. Parikh and D. Batra, "
             "“Grad-CAM: visual explanations from deep networks via gradient-based "
             "localization,” in Proc. IEEE Int. Conf. Computer Vision, 2017, pp. "
             "618–626.",

        "O. Islam, M. Assaduzzaman and M. Z. Hasan, “An explainable AI-based blood "
             "cell classification using optimized convolutional neural network,” Journal "
             "of Pathology Informatics, vol. 15, 100389, 2024.",

        "M. M. Alam and M. T. Islam, “Deep learning model for the automatic "
             "classification of white blood cells,” Computational Intelligence and "
             "Neuroscience, vol. 2022, 7384131, 2022.",

        "M. Erten, P. D. Barua, S. Dogan, T. Tuncer, R.-S. Tan and U. R. Acharya, "
             "“ConcatNeXt: an automated blood cell classification with a new deep "
             "convolutional neural network,” Multimedia Tools and Applications, vol. 84, "
             "pp. 22231–22249, 2025.",

        "M. Ghaderzadeh, F. Asadi, A. Hosseini et al., “A deep network designed for "
             "segmentation and classification of leukemia using fusion of the transfer "
             "learning models,” Complex & Intelligent Systems, vol. 7, pp. 2197–2220, "
             "2021.",

        "M. Aria, S. Meskinimood, A. Dateki, B. Nouri and A. Mohammadi, “A fast and "
             "efficient CNN model for B-ALL diagnosis and its subtypes classification "
             "using peripheral blood smear images,” International Journal of Intelligent "
             "Systems, vol. 37, no. 8, pp. 5113–5133, 2021.",

        "A. U. Rahman, S. Abbas, M. Gollapalli et al., “Lightweight EfficientNetB3 "
             "model based on depthwise separable convolutions for enhancing "
             "classification of leukemia white blood cell images,” IEEE Access, vol. 11, "
             "pp. 36490–36502, 2023.",

        "A. Acevedo, A. Merino, I. Ruiz et al., “An efficient multi-level "
             "convolutional neural network approach for white blood cells "
             "classification,” Diagnostics, vol. 12, no. 2, 248, 2022.",

        "J. Wang, “Deep learning in hematology: from molecules to patients,” "
             "Clinical Hematology International, vol. 6, no. 4, pp. 38–61, 2024.",

        "A. Rehman et al., “Automated diagnosis of leukemia: a comprehensive "
             "review,” IEEE Access, vol. 9, pp. 116156–116175, 2021.",

        "T.-Y. Lin, P. Goyal, R. Girshick, K. He and P. Dollár, “Focal loss for "
             "dense object detection,” in Proc. IEEE Int. Conf. Computer Vision, 2017, "
             "pp. 2980–2988.",

        "K. Chen, W. Lei, S. Zhao et al., “PCCT: progressive class-center triplet "
             "loss for imbalanced medical image classification,” IEEE Journal of "
             "Biomedical and Health Informatics, vol. 27, no. 4, pp. 2026–2036, 2023.",

        "N. Abbas et al., “Hybrid Inception v3 XGBoost model for acute "
             "lymphoblastic leukemia classification,” Computational and Mathematical "
             "Methods in Medicine, vol. 2021, 2577375, 2021.",

        "S. I. Khan et al., “Blood cancer prediction using leukemia microarray gene "
             "data and hybrid logistic vector trees model,” Scientific Reports, vol. 12, "
             "1882, 2022.",

        "G. An, M. Akiba, K. Omodaka, T. Nakazawa and H. Yokota, “Hierarchical deep "
             "learning models using transfer learning for disease detection and "
             "classification based on small number of medical images,” Scientific "
             "Reports, vol. 11, 4250, 2021.",

        "K. Kowsari, R. Sali, L. Ehsan et al., “HMIC: hierarchical medical image "
             "classification, a deep learning approach,” Information, vol. 11, no. 6, "
             "318, 2020.",

        "K. He, X. Zhang, S. Ren and J. Sun, “Deep residual learning for image "
             "recognition,” in Proc. IEEE Conf. Computer Vision and Pattern Recognition, "
             "2016, pp. 770–778.",

        "A. Acevedo, S. Alférez, A. Merino, L. Puigví and J. Rodellar, “Recognition "
             "of peripheral blood cell images using convolutional neural networks,” "
             "Computer Methods and Programs in Biomedicine, vol. 180, 105020, 2019.",

        "L. Wenderoth et al., “Transferable automatic hematological cell "
             "classification: overcoming data limitations with self-supervised "
             "learning,” Computer Methods and Programs in Biomedicine, vol. 260, 108560, "
             "2025.",

        "K. A. K. Al-Dulaimi et al., “Accurate classification of white blood cells "
             "by coupling pre-trained ResNet and DenseNet with SCAM mechanism,” BMC "
             "Bioinformatics, vol. 23, 329, 2022.",

        "H. Rezatofighi et al., “A deep learning framework for leukemia cancer "
             "detection in microscopic blood samples using squeeze and excitation "
             "learning,” Mathematical Problems in Engineering, vol. 2022, 2801227, 2022.",

        "S. H. Park et al., “LeuFeatx: deep learning-based feature extractor for "
             "the diagnosis of acute leukemia from microscopic images of peripheral "
             "blood smear,” Computers in Biology and Medicine, vol. 142, 105236, 2022.",

        "P. Lewicki et al., “Deep learning identifies acute promyelocytic leukemia "
             "in bone marrow smears,” BMC Cancer, vol. 22, 332, 2022.",

        "W. Wang et al., “Development and evaluation of a leukemia diagnosis system "
             "using deep learning in real clinical scenarios,” Frontiers in Pediatrics, "
             "vol. 9, 693676, 2021.",

        "E. Reinhard, M. Ashikhmin, B. Gooch and P. Shirley, “Color transfer "
             "between images,” IEEE Computer Graphics and Applications, vol. 21, no. 5, "
             "pp. 34–41, 2001.",

        "M. Macenko et al., “A method for normalizing histology slides for "
             "quantitative analysis,” in Proc. IEEE Int. Symp. Biomedical Imaging, 2009, "
             "pp. 1107–1110.",

        "A. Dosovitskiy et al., “An image is worth 16x16 words: transformers for "
             "image recognition at scale,” in Proc. Int. Conf. Learning Representations, "
             "2021.",

        "B. Kang et al., “Decoupling representation and classifier for long-tailed "
             "recognition,” in Proc. Int. Conf. Learning Representations, 2020.",

        "A. K. Menon, S. Jayasumana, A. S. Rawat, H. Jain, A. Veit and S. Kumar, "
             "“Long-tail learning via logit adjustment,” in Proc. Int. Conf. Learning "
             "Representations, 2021.",
    ]
    for i, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.45)
        p.paragraph_format.first_line_indent = Inches(-0.45)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"[{i}]\t{ref}")
        r.font.name = FONT
        r.font.size = Pt(12)
    pagebreak(doc)
